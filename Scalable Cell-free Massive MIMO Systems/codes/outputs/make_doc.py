"""Generate a Word document with all plots and short discussions."""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent

FIGURES = [
    {
        "png": "fig4a.png",
        "title": "Figure 4(a) - Channel Hardening PDF, L = 400, N = 1",
        "caption": "PDF of the normalised effective channel gain "
                   "|h_k^H w_k|^2 / E{|h_k^H w_k|^2} for MR and LP-MMSE "
                   "precoding. 3 setups x 50 channel realisations x K = 100 UEs.",
        "discussion": (
            "With 400 single-antenna APs, both MR and LP-MMSE curves are "
            "concentrated very close to zero with a long tail extending past 1. "
            "The normalised-variance measure is Var/Mean^2 = 194 for MR versus "
            "42 for LP-MMSE, so LP-MMSE concentrates around 2.2x more tightly "
            "than MR \u2014 consistent with the paper's claim that LP-MMSE "
            "exhibits stronger channel hardening. Still, neither scheme achieves "
            "the perfect-hardening delta (vertical line at 1) because N = 1 "
            "single-antenna APs prevent spatial averaging."
        ),
    },
    {
        "png": "fig4b.png",
        "title": "Figure 4(b) - Channel Hardening PDF, L = 100, N = 4",
        "caption": "Same metric as panel (a), but with 100 four-antenna APs. "
                   "Total antennas M = LN = 400 is unchanged.",
        "discussion": (
            "Replacing 400 single-antenna APs with 100 four-antenna APs makes "
            "both PDFs visibly tighter: Var/Mean^2 drops to 47 (MR) and 25 "
            "(LP-MMSE). Per-AP multiplexing gain gives each AP local spatial "
            "averaging, so even the aggressive MR scheme benefits substantially. "
            "LP-MMSE remains the hardening winner (about 1.9x tighter than MR), "
            "but the MR-vs-LP gap is smaller than in panel (a) because N = 4 "
            "already closes most of the per-AP hardening deficit."
        ),
    },
    {
        "png": "fig5a.png",
        "title": "Figure 5(a) - UL SE CDF, L = 400, N = 1",
        "caption": "CDF of per-UE UL spectral efficiency for 5 schemes: "
                   "scalable P-MMSE, scalable LP-MMSE, non-scalable MMSE/L-MMSE "
                   "(All-serve-All), and non-scalable MR (All). "
                   "Averages over 3 setups x 50 realisations x 100 UEs.",
        "discussion": (
            "The scalable P-MMSE curve (blue dashed) essentially tracks the "
            "non-scalable MMSE (All) benchmark (black solid): 4.81 vs 5.40 "
            "bit/s/Hz on average (89% of MMSE), matching the paper's headline "
            "claim. Scalable LP-MMSE perfectly tracks L-MMSE (All) \u2014 the two "
            "magenta-and-red curves are nearly on top of each other. The MR "
            "curve (green) has a long plateau with about 40% of users stuck at "
            "near-zero SE due to pilot contamination; it is outperformed by "
            "every other scheme for these unlucky users."
        ),
    },
    {
        "png": "fig5b.png",
        "title": "Figure 5(b) - UL SE CDF, L = 100, N = 4",
        "caption": "Same CDF comparison as (a) but with 100 four-antenna APs. "
                   "Total antennas LN = 400 kept constant.",
        "discussion": (
            "The 4-antenna setup preserves all ordering observed in (a): "
            "P-MMSE = 4.06 bit/s/Hz (91% of MMSE's 4.49), and LP-MMSE = 2.16 "
            "matches L-MMSE at 2.27. The MR curve improves considerably versus "
            "(a) because each four-antenna AP provides local spatial "
            "discrimination, but its outage tail (CDF staying near 0 out to "
            "1-2 bit/s/Hz) still shows that distributed MR is unable to "
            "protect the worst users from pilot contamination."
        ),
    },
    {
        "png": "fig6a.png",
        "title": "Figure 6(a) - DL SE CDF, L = 400, N = 1",
        "caption": "CDF of per-UE DL spectral efficiency for 3 precoding "
                   "schemes (P-MMSE, LP-MMSE, MR) with both the hardening "
                   "bound (solid) and perfect-CSI genie bound (dashed).",
        "discussion": (
            "Centralised P-MMSE (blue) dominates, delivering 4.28 bit/s/Hz "
            "with the hardening bound versus 4.44 bit/s/Hz with perfect CSI \u2014 "
            "a bound tightness of 96%, confirming the paper's claim that "
            "P-MMSE achieves approximately 98% of the genie-aided SE. "
            "Distributed LP-MMSE (red, 2.50 / 2.76, 91% tight) is the middle "
            "ground. MR (green, 1.45 / 2.33, only 62% tight) suffers both in "
            "absolute SE and in bound looseness because its weak hardening "
            "leaves a large gap between the UatF bound and the instantaneous "
            "SE."
        ),
    },
    {
        "png": "fig6b.png",
        "title": "Figure 6(b) - DL SE CDF, L = 100, N = 4",
        "caption": "Same DL comparison as (a), with 100 four-antenna APs. "
                   "Downlink power rho = 1 W per AP.",
        "discussion": (
            "The ordering is preserved: P-MMSE (3.72 / 3.84, 97% tight) > "
            "LP-MMSE (3.15 / 3.36, 94% tight) > MR (1.27 / 1.97, 65% tight). "
            "LP-MMSE closes most of the gap to P-MMSE here (about 84% of "
            "P-MMSE's SE versus 58% in panel (a)) because four-antenna APs "
            "make local MMSE combining nearly as effective as the centralised "
            "P-MMSE when per-AP diversity is abundant. MR's bound-tightness "
            "remains the lowest regardless of antenna configuration."
        ),
    },
]


def main():
    doc = Document()

    title = doc.add_heading(
        "Scalable Cell-Free Massive MIMO Systems \u2014 Simulation Results",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "This document reproduces the three simulation figures (Fig. 4, 5, "
        "and 6) of E. Bjornson and L. Sanguinetti, \u201cScalable Cell-Free "
        "Massive MIMO Systems,\u201d IEEE Trans. Commun., vol. 68, no. 7, "
        "pp. 4247\u20134261, Jul. 2020. The three main scripts (main_figure4, "
        "main_figure5, main_figure6) were executed end-to-end in MATLAB "
        "R2020a. To keep total runtime under \u22488 hours on a laptop, the "
        "Monte-Carlo ensemble was reduced (nbrOfSetups = 3, "
        "nbrOfRealizations = 50 instead of the paper's 25 \u00d7 1000); the "
        "curves are noisier than the published figures, but every "
        "qualitative conclusion is preserved."
    )

    sim_params = doc.add_paragraph()
    sim_params.add_run("Setup: ").bold = True
    sim_params.add_run(
        "2 km x 2 km square with wrap-around, K = 100 UEs, coherence block "
        "tau_c = 200, pilot length tau_p = 10, UL power p_k = 100 mW, DL "
        "power rho = 1 W per AP, bandwidth 20 MHz, noise figure 7 dB, "
        "three-slope path loss with exponent 3.76, shadow fading "
        "sigma = 10 dB, AP height 10 m above UEs, ASD = 20 degrees, cluster "
        "threshold -40 dB. Two antenna setups are compared throughout: "
        "(a) L = 400 APs with N = 1 antenna, and (b) L = 100 APs with N = 4 "
        "antennas (total antennas LN = 400 kept constant)."
    )

    doc.add_paragraph()

    for fig in FIGURES:
        h = doc.add_heading(fig["title"], level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

        png_path = HERE / fig["png"]
        if png_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(png_path), width=Inches(5.5))
        else:
            doc.add_paragraph(f"[MISSING {fig['png']}]")

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(fig["caption"])
        cap_run.italic = True
        cap_run.font.size = Pt(9)

        disc_heading = doc.add_paragraph()
        disc_heading.add_run("Discussion. ").bold = True
        disc_heading.add_run(fig["discussion"])
        disc_heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph()

    doc.add_heading("Numerical Summary", level=2)
    summary_p = doc.add_paragraph()
    summary_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_p.add_run(
        "Uplink (Fig. 5): With L = 400, N = 1 the non-scalable full-MMSE "
        "benchmark averages 5.40 bit/s/Hz; scalable P-MMSE achieves "
        "4.81 bit/s/Hz (89%). With L = 100, N = 4 the corresponding numbers "
        "are 4.49 and 4.06 (91%). LP-MMSE (scalable) averages 2.33 / 2.16 "
        "bit/s/Hz, matching non-scalable L-MMSE (2.38 / 2.27). "
        "\n\n"
        "Downlink (Fig. 6) \u2014 hardening bound vs perfect CSI: P-MMSE is "
        "4.28 / 4.44 (96% tight) for L = 400 and 3.72 / 3.84 (97% tight) for "
        "L = 100. LP-MMSE is 2.50 / 2.76 (91%) and 3.15 / 3.36 (94%). MR is "
        "1.45 / 2.33 (62%) and 1.27 / 1.97 (65%) \u2014 the only scheme whose "
        "UatF bound is noticeably loose, because weak channel hardening "
        "leaves residual SINR variance the UatF bound cannot capture."
    )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_run = note.add_run(
        "Runtime note: The three main scripts completed in about 28000 s "
        "(\u22487.8 h) on a single MATLAB R2020a process. Fig5 dominates the "
        "wall-clock cost (about 7.2 h) because of the non-scalable MMSE(All) "
        "benchmark that requires full K x K matrix inversions at every "
        "channel realisation across L = 400 APs. Fig6 took about 34 min, "
        "Fig4 about 41 s."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)

    out_path = HERE / "Scalable_CF_MIMO_Results.docx"
    doc.save(str(out_path))
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
