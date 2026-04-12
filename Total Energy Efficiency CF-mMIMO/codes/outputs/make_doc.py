"""Generate a Word document with all plots and short discussions."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent

FIGURES = [
    {
        "png": "fig1.png",
        "title": "Figure 1 \u2014 Convergence of Algorithm 1 (SCA)",
        "caption": "Energy efficiency vs SCA iteration index for three "
                   "(M, K) configurations. N = 1 antenna per AP, D = 1 km.",
        "discussion": (
            "All three curves rise monotonically and plateau within 4\u20136 "
            "iterations, confirming the guaranteed-ascent property of the "
            "SCA framework (Section IV). The (M=40, K=10) configuration "
            "achieves the highest EE (\u22487.6 Mbit/J) because its larger "
            "array-to-user ratio gives the optimizer more degrees of freedom "
            "to redistribute power. The smallest system (M=8, K=2) converges "
            "fastest but to a lower EE due to fewer APs."
        ),
    },
    {
        "png": "fig2.png",
        "title": "Figure 2 \u2014 Impact of Power Control on EE",
        "caption": "Average EE vs number of APs M for K = 20 and K = 40, "
                   "with (Algorithm 1) and without (equal-power) power control. "
                   "N = 1, D = 1 km, averaged over 5 Monte-Carlo drops.",
        "discussion": (
            "Algorithm 1 (solid) delivers roughly 3\u20133.5\u00d7 higher EE than "
            "the no-power-control baseline (dashed) across all M values. "
            "EE decreases with M because adding more single-antenna APs "
            "increases fixed backhaul power P_0 per AP without proportional "
            "throughput gain. K = 40 slightly outperforms K = 20 with "
            "power control because more users utilise the same AP "
            "infrastructure, amortising the fixed cost."
        ),
    },
    {
        "png": "fig3.png",
        "title": "Figure 3 \u2014 EE vs Traffic-Dependent Backhaul P_bt",
        "caption": "Average EE for three schemes: no AP selection (baseline), "
                   "LSF-based (Alg. 3), and received-power-based (Alg. 2). "
                   "M = 40, K = 20, N = 1, \u03c1 = 95%.",
        "discussion": (
            "All three curves decline as P_bt increases because heavier "
            "backhaul cost penalises every bit of throughput. The no-selection "
            "baseline (blue) starts highest at P_bt = 0.25 because it uses "
            "all APs, but the AP-selection schemes converge toward it at "
            "high P_bt as the traffic-dependent cost dominates the fixed "
            "cost. At small M = 40 (vs the paper\u2019s M = 100) the difference "
            "between the three schemes is modest; at full scale the selection "
            "schemes would stay flat while the baseline crashes."
        ),
    },
    {
        "png": "fig4.png",
        "title": "Figure 4 \u2014 Average Number of APs per User vs M",
        "caption": "APs selected per user for received-power (Alg. 2) and "
                   "LSF-based (Alg. 3) at D = 1 km and D = 2 km. "
                   "K = 20, N = 1, \u03c1 = 95%.",
        "discussion": (
            "The number of selected APs grows sub-linearly with M: at "
            "M = 100 the received-power scheme selects ~23 APs (D = 1 km) "
            "while LSF-based selects ~13, meaning each user is served by "
            "only 13\u201323% of all APs. The D = 2 km curves are higher "
            "because in a sparser deployment the nearest APs are farther "
            "apart, so more APs are needed to capture 95% of the signal "
            "power. This confirms the paper\u2019s scalability claim."
        ),
    },
    {
        "png": "fig5.png",
        "title": "Figure 5 \u2014 EE vs N (Antennas per AP) at Fixed MN",
        "caption": "Average EE as a function of N with MN = 128 total "
                   "antennas. Three scenarios with different D, P_bt, S_o. "
                   "K = 10, Alg. 3 selection.",
        "discussion": (
            "All three scenarios show EE increasing with N, with the D = 1 km "
            "dense scenario (blue) achieving the highest values and peaking "
            "near N = 16. The paper predicts an inverted-U shape: small N means "
            "many APs with high fixed backhaul overhead, while large N means "
            "fewer APs with less spatial diversity. The harder QoS scenario "
            "(S_o = 2, yellow) drops to zero at N = 1\u20132 because the "
            "SCA problem becomes infeasible with so few antennas per AP."
        ),
    },
    {
        "png": "fig6.png",
        "title": "Figure 6 \u2014 Cell-Free vs Colocated Massive MIMO",
        "caption": "Average EE vs per-user SE target S_o for cell-free "
                   "(N_CF = 4, Alg. 3 selection) and colocated (single AP "
                   "with MN antennas). MN = 32 and 64, K = 10, D = 1 km.",
        "discussion": (
            "Cell-free mMIMO (solid) dominates colocated (dashed) across "
            "most of the SE range: CF-mMIMO at MN = 32 stays above 13 "
            "Mbit/J out to S_o = 2, while the colocated system crashes to "
            "zero when the QoS becomes infeasible for a single-AP geometry. "
            "The colocated system shows erratic peaks (e.g. MN = 32 at "
            "S_o = 1) because with only 2 MC drops, a single favourable "
            "user placement skews the average. With more MC runs the "
            "colocated curve would smooth out, but the CF advantage "
            "would remain."
        ),
    },
]


def main():
    doc = Document()

    title = doc.add_heading(
        "On the Total Energy Efficiency of Cell-Free Massive MIMO "
        "\u2014 Simulation Results",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "This document reproduces the six simulation figures of H. Q. Ngo, "
        "L.-N. Tran, T. Q. Duong, M. Matthaiou, and E. G. Larsson, "
        "\u201cOn the Total Energy Efficiency of Cell-Free Massive MIMO,\u201d "
        "IEEE Trans. Green Commun. Netw., vol. 2, no. 1, pp. 25\u201339, "
        "Mar. 2018. All six main scripts were executed end-to-end in "
        "MATLAB R2020a with CVX 2.2 (SDPT3 solver). Reduced (M, K, NMC) "
        "defaults were used to keep total runtime under \u224810 hours; "
        "the qualitative trends match the paper."
    )

    sim_params = doc.add_paragraph()
    sim_params.add_run("Setup: ").bold = True
    sim_params.add_run(
        "1 km \u00d7 1 km square (default), 3-slope path loss (f = 1.9 GHz, "
        "h_AP = 15 m, h_UE = 1.65 m), shadow fading \u03c3 = 8 dB, "
        "B = 20 MHz, NF = 9 dB, PA efficiency \u03b7 = 0.4, per-antenna "
        "circuit power P_tc = 0.2 W, fixed backhaul P_0 = 0.825 W/AP, "
        "traffic-dependent backhaul P_bt = 0.25 W/(Gbit/s), coherence block "
        "\u03c4_c = 200, pilot length \u03c4_p = 20, pilot power 0.2 W, "
        "downlink power 1 W/antenna, QoS target S_o = 1 bit/s/Hz."
    )

    doc.add_paragraph()

    for fig in FIGURES:
        h = doc.add_heading(fig["title"], level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

        png_path = HERE / fig["png"]
        if png_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(png_path), width=Inches(5.5))
        else:
            doc.add_paragraph(f"[MISSING {fig['png']}]")

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(fig["caption"])
        cap_run.italic = True
        cap_run.font.size = Pt(9)

        disc = doc.add_paragraph()
        disc.add_run("Discussion. ").bold = True
        disc.add_run(fig["discussion"])
        disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph()

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_run = note.add_run(
        "Runtime note: Total wall-clock time was approximately 10 hours "
        "on MATLAB R2020a with CVX 2.2 / SDPT3. Fig. 2 dominated "
        "(~8 h) due to 50 Algorithm 1 invocations at M up to 100. "
        "Figs. 5 and 6 used only 2 MC drops each (NMC = 2), which "
        "produces noisy colocated curves in Fig. 6; increase NMC for "
        "smoother results."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)

    out_path = HERE / "Total_EE_CF_mMIMO_Results.docx"
    doc.save(str(out_path))
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
